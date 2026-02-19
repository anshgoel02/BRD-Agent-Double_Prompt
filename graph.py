from __future__ import annotations

import json
from pathlib import Path
import re
from typing import Any, List

from langgraph.graph import END, StateGraph

from .config import GENERATOR_PROMPT, INTAKE_GAP_PROMPT, TEMPLATE_SECTIONS
from .debug import debug_state
from .file_loaders import is_image_file, iter_input_files, load_source_text
from .llm_utils import generate_text_with_usage, parse_json_obj
from .models import BRDState, IntakeGapResult, SectionCoverage

try:
    from docx import Document
except Exception:
    Document = None


def _ensure_sources(state: BRDState) -> BRDState:
    files = list(iter_input_files(state.inputs))
    texts: List[str] = []
    images: List[str] = []
    for file_path in files:
        if is_image_file(file_path):
            images.append(str(file_path))
            continue
        try:
            text = load_source_text(file_path)
        except Exception:
            continue
        if text.strip():
            texts.append(f"[SOURCE: {file_path.name}]\n{text}")
    state.source_texts = texts
    state.image_paths = images
    return state


def _validate_intake_payload(obj: dict) -> tuple[bool, str]:
    required_top_keys = {"FACT_PACK", "TEMPLATE_COVERAGE", "GAP_REPORT", "QUICK_SUMMARY"}
    missing_top = [k for k in required_top_keys if k not in obj]
    if missing_top:
        return False, f"Missing top-level keys: {', '.join(missing_top)}"

    if not isinstance(obj.get("FACT_PACK"), list):
        return False, "FACT_PACK must be a list."
    if not isinstance(obj.get("TEMPLATE_COVERAGE"), list):
        return False, "TEMPLATE_COVERAGE must be a list."
    if not isinstance(obj.get("GAP_REPORT"), dict):
        return False, "GAP_REPORT must be an object."
    if not isinstance(obj.get("QUICK_SUMMARY"), list):
        return False, "QUICK_SUMMARY must be a list."

    for idx, item in enumerate(obj["FACT_PACK"], start=1):
        if not isinstance(item, dict):
            return False, f"FACT_PACK[{idx}] must be an object."
        required = {"id", "category", "fact", "source"}
        missing = [k for k in required if not str(item.get(k, "")).strip()]
        if missing:
            return False, f"FACT_PACK[{idx}] missing required fields: {', '.join(missing)}"

    allowed_status = {"covered", "partially covered", "not covered"}
    for idx, item in enumerate(obj["TEMPLATE_COVERAGE"], start=1):
        if not isinstance(item, dict):
            return False, f"TEMPLATE_COVERAGE[{idx}] must be an object."
        required = {"section_index", "section_name", "status", "why"}
        missing = [k for k in required if not str(item.get(k, "")).strip()]
        if missing:
            return False, f"TEMPLATE_COVERAGE[{idx}] missing required fields: {', '.join(missing)}"
        status = str(item.get("status", "")).strip().lower()
        if status not in allowed_status:
            return False, (
                f"TEMPLATE_COVERAGE[{idx}].status must be one of: "
                "Covered | Partially Covered | Not Covered"
            )

    gap = obj["GAP_REPORT"]
    required_buckets = {"blocking", "important", "nice_to_have"}
    missing_buckets = [k for k in required_buckets if k not in gap]
    if missing_buckets:
        return False, f"GAP_REPORT missing buckets: {', '.join(missing_buckets)}"

    allowed_priority = {"blocking", "important", "nice-to-have"}
    for bucket in ("blocking", "important", "nice_to_have"):
        items = gap.get(bucket)
        if not isinstance(items, list):
            return False, f"GAP_REPORT.{bucket} must be a list."
        for idx, item in enumerate(items, start=1):
            if not isinstance(item, dict):
                return False, f"GAP_REPORT.{bucket}[{idx}] must be an object."
            required = {"id", "template_section", "priority", "missing", "ask_back"}
            missing = [k for k in required if not str(item.get(k, "")).strip()]
            if missing:
                return False, f"GAP_REPORT.{bucket}[{idx}] missing required fields: {', '.join(missing)}"
            priority = str(item.get("priority", "")).strip().lower()
            if priority not in allowed_priority:
                return False, (
                    f"GAP_REPORT.{bucket}[{idx}].priority must be one of: "
                    "Blocking | Important | Nice-to-have"
                )

    return True, ""


def intake_gap_node(state: BRDState) -> BRDState:
    print("[node] intake_gap")
    state = _ensure_sources(state)
    prompt = INTAKE_GAP_PROMPT.format(
        TEMPLATE_SECTIONS="\n".join(f"- {s}" for s in TEMPLATE_SECTIONS),
        INPUTS_TEXT="\n\n".join(state.source_texts),
    )
    raw, usage = generate_text_with_usage(prompt, max_tokens=100000, image_paths=state.image_paths)
    obj = parse_json_obj(raw)
    is_valid, error = _validate_intake_payload(obj) if obj else (False, "Empty or invalid JSON.")
    if not is_valid:
        print(f"[warn] intake_gap payload invalid ({error}); retrying once")
        retry_prompt = (
            "Return valid JSON only and include all required keys/fields exactly as schema.\n\n"
            + prompt
        )
        raw, usage = generate_text_with_usage(retry_prompt, max_tokens=100000, image_paths=state.image_paths)
        obj = parse_json_obj(raw)
        is_valid, error = _validate_intake_payload(obj) if obj else (False, "Empty or invalid JSON.")
    if not is_valid:
        raise ValueError(
            "Intake payload validation failed after retry. "
            f"Reason: {error}. "
            "Expected top-level keys: FACT_PACK, TEMPLATE_COVERAGE, GAP_REPORT, QUICK_SUMMARY."
        )

    fact_pack = obj.get("FACT_PACK", [])
    template_coverage = obj.get("TEMPLATE_COVERAGE", [])
    gap_report = obj.get("GAP_REPORT", {})
    quick_summary = obj.get("QUICK_SUMMARY", [])

    key_facts: list[str] = []
    if isinstance(fact_pack, list):
        for item in fact_pack:
            if not isinstance(item, dict):
                continue
            category = str(item.get("category", "")).strip()
            fact = str(item.get("fact", "")).strip()
            source = str(item.get("source", "")).strip()
            if fact:
                if category and source:
                    key_facts.append(f"{category}: {fact} (Source: {source})")
                elif category:
                    key_facts.append(f"{category}: {fact}")
                else:
                    key_facts.append(fact)

    coverage: list[SectionCoverage] = []
    if isinstance(template_coverage, list):
        for item in template_coverage:
            if not isinstance(item, dict):
                continue
            section = str(item.get("section_name", "")).strip() or str(item.get("section", "")).strip()
            raw_status = str(item.get("status", "")).strip().lower()
            why = str(item.get("why", "")).strip()
            if not section:
                continue
            if "partially" in raw_status:
                mapped_status = "partial"
            elif "not covered" in raw_status or "missing" in raw_status:
                mapped_status = "missing"
            elif "covered" in raw_status:
                mapped_status = "covered"
            else:
                mapped_status = "missing"
            coverage.append(SectionCoverage(section=section, status=mapped_status, notes=why))

    missing_information: list[str] = []
    open_questions: list[str] = []
    if isinstance(gap_report, dict):
        for priority in ("blocking", "important", "nice_to_have"):
            items = gap_report.get(priority, [])
            if not isinstance(items, list):
                continue
            for item in items:
                if not isinstance(item, dict):
                    continue
                section = str(item.get("template_section", "")).strip()
                missing = str(item.get("missing", "")).strip()
                ask_back = str(item.get("ask_back", "")).strip()
                if missing:
                    prefix = f"[{priority}] "
                    if section:
                        missing_information.append(f"{prefix}{section}: {missing}")
                    else:
                        missing_information.append(f"{prefix}{missing}")
                if ask_back:
                    open_questions.append(ask_back)

    summary_lines = [str(x).strip() for x in quick_summary if str(x).strip()] if isinstance(quick_summary, list) else []
    project_summary = " ".join(summary_lines[:6]).strip()

    if not coverage:
        for section in TEMPLATE_SECTIONS:
            coverage.append(SectionCoverage(section=section, status="missing", notes="Not assessed"))

    state.intake_gap = IntakeGapResult(
        project_summary=project_summary,
        key_facts=key_facts,
        template_coverage=coverage,
        missing_information=missing_information,
        open_questions=open_questions,
    )
    if state.output_intake_docx_path:
        _save_intake_gap_docx(state.intake_gap, state.output_intake_docx_path)
    print(
        f"[node] intake_gap tokens_used={usage['total_tokens']} "
        f"(prompt={usage['prompt_tokens']} completion={usage['completion_tokens']})"
    )
    debug_state("intake_gap", state)
    return state


def intake_review_node(state: BRDState) -> BRDState:
    print("[node] intake_review")
    print("\n--- INTAKE SUMMARY ---\n")
    print(state.intake_gap.project_summary or "No summary generated.")
    print("\n--- MISSING INFORMATION ---\n")
    if state.intake_gap.missing_information:
        for item in state.intake_gap.missing_information:
            print(f"- {item}")
    else:
        print("No missing information detected.")

    choice = input("Approve intake and proceed to BRD generation? (y/n): ").strip().lower()
    if choice in {"y", "yes"}:
        state.intake_approved = True
    else:
        state.intake_approved = False
        extra = input("Provide additional input paths (comma-separated), or leave empty: ").strip()
        if extra:
            for item in [p.strip() for p in extra.split(",") if p.strip()]:
                state.inputs.append(item)
    print("[node] intake_review tokens_used=0")
    debug_state("intake_review", state)
    return state


def generator_node(state: BRDState) -> BRDState:
    print("[node] generator")
    if not state.source_texts:
        state = _ensure_sources(state)

    base_prompt = GENERATOR_PROMPT.format(
        PROJECT_SUMMARY=state.intake_gap.project_summary or "TBC",
        KEY_FACTS="\n".join(f"- {x}" for x in state.intake_gap.key_facts) or "- TBC",
        MISSING_INFORMATION="\n".join(f"- {x}" for x in state.intake_gap.missing_information) or "- None",
        OPEN_QUESTIONS="\n".join(f"- {x}" for x in state.intake_gap.open_questions) or "- None",
        HUMAN_FEEDBACK=state.generation_feedback or "None",
        INPUTS_TEXT="\n\n".join(state.source_texts),
    )
    prompt = _document_schema_prompt(base_prompt)
    raw, usage = generate_text_with_usage(prompt, max_tokens=100000, image_paths=state.image_paths)
    spec = _extract_doc_spec(raw)
    if spec is None:
        print("[warn] generator invalid/empty JSON spec; retrying once")
        retry_prompt = _document_schema_prompt("Return valid JSON only.\n\n" + base_prompt)
        raw, usage = generate_text_with_usage(retry_prompt, max_tokens=100000, image_paths=state.image_paths)
        spec = _extract_doc_spec(raw)
    if spec is None:
        raise ValueError("Generator failed to return structured document schema after retry.")

    markdown = _doc_spec_to_markdown(spec)
    # ok, issues = _quality_gate(markdown, bool(state.intake_gap.missing_information))
    # if not ok:
    #     print("[warn] generator quality gate failed; attempting one corrective regeneration")
    #     fix_prompt = _document_schema_prompt(
    #         base_prompt
    #         + "\n\nQuality issues to fix:\n"
    #         + "\n".join(f"- {issue}" for issue in issues)
    #         + "\n\nReturn corrected JSON schema only."
    #     )
    #     raw, usage = generate_text_with_usage(fix_prompt, max_tokens=100000, image_paths=state.image_paths)
    #     spec = _extract_doc_spec(raw)
    #     if spec is None:
    #         raise ValueError("Corrective generation failed: invalid structured document schema.")
    #     markdown = _doc_spec_to_markdown(spec)
    #     ok, issues = _quality_gate(markdown, bool(state.intake_gap.missing_information))
    #     if not ok:
    #         raise ValueError("Quality gate failed after corrective regeneration: " + "; ".join(issues))

    state.brd_markdown = markdown.strip()
    _save_revision_history(state, state.brd_markdown, spec)

    if state.output_markdown_path:
        Path(state.output_markdown_path).write_text(state.brd_markdown, encoding="utf-8")
    if state.output_docx_path:
        _render_docx_from_spec(spec, state.output_docx_path)

    print(
        f"[node] generator tokens_used={usage['total_tokens']} "
        f"(prompt={usage['prompt_tokens']} completion={usage['completion_tokens']})"
    )
    debug_state("generator", state)
    return state


def generation_review_node(state: BRDState) -> BRDState:
    print("[node] generation_review")
    print("\n--- GENERATED BRD ---\n")
    print(state.brd_markdown)
    choice = input("Approve generated BRD? (y/n): ").strip().lower()
    if choice in {"y", "yes"}:
        state.generation_approved = True
        state.generation_feedback = None
    else:
        state.generation_approved = False
        state.generation_feedback = input("Provide revision suggestions: ").strip()
    print("[node] generation_review tokens_used=0")
    debug_state("generation_review", state)
    return state


def _save_docx(markdown: str, output_path: str) -> None:
    if Document is None:
        raise RuntimeError("python-docx is required to write .docx output.")
    doc = Document()
    for line in markdown.splitlines():
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)
    doc.save(output_path)


def _document_schema_prompt(base_prompt: str) -> str:
    return f"""{base_prompt}

Return the BRD as JSON only. Do not include markdown fences, prose, or explanations.
Use this exact schema:
{{
  "title": "string",
  "blocks": [
    {{
      "type": "heading",
      "level": 1,
      "text": "string"
    }},
    {{
      "type": "paragraph",
      "text": "string"
    }},
    {{
      "type": "bullet_list",
      "items": [
        {{"text": "string", "level": 0}},
        {{"text": "string", "level": 1}}
      ]
    }},
    {{
      "type": "numbered_list",
      "items": [
        {{"text": "string", "level": 0}},
        {{"text": "string", "level": 1}}
      ]
    }},
    {{
      "type": "table",
      "headers": ["col1", "col2"],
      "rows": [["r1c1", "r1c2"], ["r2c1", "r2c2"]]
    }}
  ]
}}
"""


def _extract_doc_spec(output: str) -> dict[str, Any] | None:
    text = output.strip()
    if not text:
        return None
    try:
        parsed = json.loads(text)
    except json.JSONDecodeError:
        start = text.find("{")
        end = text.rfind("}")
        if start == -1 or end == -1 or end <= start:
            return None
        try:
            parsed = json.loads(text[start : end + 1])
        except json.JSONDecodeError:
            return None
    if not isinstance(parsed, dict):
        return None
    if not isinstance(parsed.get("blocks"), list):
        return None
    return parsed


def _doc_spec_to_markdown(spec: dict[str, Any]) -> str:
    lines: list[str] = []
    title = str(spec.get("title", "")).strip()
    if title:
        lines.append(f"# {title}")
        lines.append("")

    for block in spec.get("blocks", []):
        if not isinstance(block, dict):
            continue
        btype = block.get("type")
        if btype == "heading":
            level = int(block.get("level", 2))
            level = min(max(level, 1), 6)
            text = str(block.get("text", "")).strip()
            if text:
                lines.append(f"{'#' * level} {text}")
                lines.append("")
        elif btype == "paragraph":
            text = str(block.get("text", "")).strip()
            if text:
                lines.append(text)
                lines.append("")
        elif btype in {"bullet_list", "numbered_list"}:
            items = block.get("items", [])
            if isinstance(items, list):
                idx = 1
                for item in items:
                    if isinstance(item, dict):
                        text = str(item.get("text", "")).strip()
                        level = int(item.get("level", 0))
                    else:
                        text = str(item).strip()
                        level = 0
                    if not text:
                        continue
                    indent = "  " * max(level, 0)
                    if btype == "numbered_list":
                        lines.append(f"{indent}{idx}. {text}")
                        idx += 1
                    else:
                        lines.append(f"{indent}- {text}")
                lines.append("")
        elif btype == "table":
            headers = block.get("headers", [])
            rows = block.get("rows", [])
            if isinstance(headers, list) and headers:
                header_cells = [str(h) for h in headers]
                lines.append("| " + " | ".join(header_cells) + " |")
                lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
                if isinstance(rows, list):
                    for row in rows:
                        if isinstance(row, list):
                            cells = [str(c) for c in row]
                            if len(cells) < len(header_cells):
                                cells.extend([""] * (len(header_cells) - len(cells)))
                            lines.append("| " + " | ".join(cells[: len(header_cells)]) + " |")
                lines.append("")
    return "\n".join(lines).strip()


def _render_docx_from_spec(spec: dict[str, Any], output_path: str) -> None:
    if Document is None:
        raise RuntimeError("python-docx is required to write .docx output.")

    doc = Document()
    title = str(spec.get("title", "")).strip()
    if title:
        doc.add_heading(title, level=1)

    for block in spec.get("blocks", []):
        if not isinstance(block, dict):
            continue
        btype = block.get("type")
        if btype == "heading":
            level = int(block.get("level", 2))
            level = min(max(level, 1), 4)
            text = str(block.get("text", "")).strip()
            if text:
                doc.add_heading(text, level=level)
        elif btype == "paragraph":
            text = str(block.get("text", "")).strip()
            if text:
                doc.add_paragraph(text)
        elif btype in {"bullet_list", "numbered_list"}:
            items = block.get("items", [])
            if isinstance(items, list):
                for item in items:
                    if isinstance(item, dict):
                        text = str(item.get("text", "")).strip()
                        level = max(int(item.get("level", 0)), 0)
                    else:
                        text = str(item).strip()
                        level = 0
                    if not text:
                        continue
                    if btype == "numbered_list":
                        style = "List Number" if level <= 0 else f"List Number {min(level + 1, 3)}"
                    else:
                        style = "List Bullet" if level <= 0 else f"List Bullet {min(level + 1, 3)}"
                    try:
                        doc.add_paragraph(text, style=style)
                    except KeyError:
                        fallback = "List Number" if btype == "numbered_list" else "List Bullet"
                        try:
                            doc.add_paragraph(text, style=fallback)
                        except KeyError:
                            doc.add_paragraph(text)
        elif btype == "table":
            headers = block.get("headers", [])
            rows = block.get("rows", [])
            if isinstance(headers, list) and headers:
                cols = len(headers)
                table = doc.add_table(rows=1, cols=cols)
                table.style = "Table Grid"
                for i, h in enumerate(headers):
                    table.rows[0].cells[i].text = str(h)
                if isinstance(rows, list):
                    for row in rows:
                        if not isinstance(row, list):
                            continue
                        cells = table.add_row().cells
                        for i in range(cols):
                            cells[i].text = str(row[i] if i < len(row) else "")
    doc.save(output_path)


def _quality_gate(markdown: str, has_missing_info: bool) -> tuple[bool, list[str]]:
    issues: list[str] = []
    lower = markdown.lower()

    required_sections = [
        "executive summary",
        "business context & problem statement",
        "objectives & success metrics",
        "scope",
        "stakeholders & roles",
        "functional requirements",
        "non-functional requirements",
        "integrations & interfaces",
        "data & fields",
        "slas & kpis",
        "risks, assumptions & dependencies",
        "timeline & milestones",
        "open items / decisions needed",
    ]
    for section in required_sections:
        if section not in lower:
            issues.append(f"Missing required section heading/content: {section}")

    if re.search(r"\bFR-\d+\b", markdown) is None:
        issues.append("Missing functional requirement IDs (FR-1, FR-2, ...).")
    if re.search(r"\bNFR-\d+\b", markdown) is None:
        issues.append("Missing non-functional requirement IDs (NFR-1, NFR-2, ...).")

    if has_missing_info:
        if "open items / decisions needed" not in lower and "open questions" not in lower:
            issues.append("Missing Open Questions/Open Items section while missing information exists.")
        else:
            has_question_bullets = bool(re.search(r"(?im)^(?:-|\d+\.)\s+", markdown))
            if not has_question_bullets:
                issues.append("Open Questions/Open Items section appears empty while missing information exists.")

    return (len(issues) == 0, issues)


def _save_revision_history(state: BRDState, markdown: str, spec: dict[str, Any] | None) -> None:
    history_dir = Path(state.revision_history_dir or ".")
    history_dir.mkdir(parents=True, exist_ok=True)

    base_name = "brd_double"
    if state.output_markdown_path:
        base_name = Path(state.output_markdown_path).stem or base_name

    state.revision_number += 1
    version = state.revision_number
    md_path = history_dir / f"{base_name}_v{version}.md"
    docx_path = history_dir / f"{base_name}_v{version}.docx"

    md_path.write_text(markdown, encoding="utf-8")
    if spec is not None:
        _render_docx_from_spec(spec, str(docx_path))
    else:
        _save_docx(markdown, str(docx_path))

    state.revision_files.append(str(md_path))
    state.revision_files.append(str(docx_path))


def _save_intake_gap_docx(intake: IntakeGapResult, output_path: str) -> None:
    if Document is None:
        raise RuntimeError("python-docx is required to write .docx output.")

    doc = Document()
    doc.add_heading("Intake Gap Result", level=1)

    doc.add_heading("Project Summary", level=2)
    doc.add_paragraph(intake.project_summary or "No summary generated.")

    doc.add_heading("Key Facts", level=2)
    if intake.key_facts:
        for fact in intake.key_facts:
            doc.add_paragraph(str(fact), style="List Bullet")
    else:
        doc.add_paragraph("None")

    doc.add_heading("Template Coverage", level=2)
    if intake.template_coverage:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Section"
        table.rows[0].cells[1].text = "Status"
        table.rows[0].cells[2].text = "Notes"
        for item in intake.template_coverage:
            row = table.add_row().cells
            row[0].text = item.section
            row[1].text = item.status
            row[2].text = item.notes
    else:
        doc.add_paragraph("No coverage data.")

    doc.add_heading("Missing Information", level=2)
    if intake.missing_information:
        for gap in intake.missing_information:
            doc.add_paragraph(str(gap), style="List Bullet")
    else:
        doc.add_paragraph("None")

    doc.add_heading("Open Questions", level=2)
    if intake.open_questions:
        for q in intake.open_questions:
            doc.add_paragraph(str(q), style="List Number")
    else:
        doc.add_paragraph("None")

    doc.save(output_path)


def build_graph() -> StateGraph:
    graph = StateGraph(BRDState)
    graph.add_node("intake_gap", intake_gap_node)
    graph.add_node("intake_review", intake_review_node)
    graph.add_node("generator", generator_node)
    graph.add_node("generation_review", generation_review_node)

    graph.set_entry_point("intake_gap")
    graph.add_edge("intake_gap", "intake_review")

    graph.add_conditional_edges(
        "intake_review",
        lambda s: "generator" if s.intake_approved else "intake_gap",
        {"generator": "generator", "intake_gap": "intake_gap"},
    )

    graph.add_edge("generator", "generation_review")
    graph.add_conditional_edges(
        "generation_review",
        lambda s: END if s.generation_approved else "generator",
        {END: END, "generator": "generator"},
    )
    return graph
